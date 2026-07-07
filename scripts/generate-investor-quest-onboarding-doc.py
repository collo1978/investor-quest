#!/usr/bin/env python3
"""Generate Investor Quest developer onboarding .docx."""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Installing python-docx...", file=sys.stderr)
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Investor-Quest-Developer-Onboarding.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build() -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("Investor Quest", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Developer Onboarding & Product Vision")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    doc.add_paragraph("Document purpose: onboard a developer to build a working prototype modeled after this project.")
    doc.add_paragraph("Version: prototype reference · NVIDIA 10-K Schools demo")
    doc.add_paragraph()

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "Investor Quest is a premium, gamified stock-research learning platform. It transforms "
        "corporate annual reports (10-K filings) into an interactive adventure that helps students "
        "— especially middle and high school learners — build real investor literacy without reading "
        "80,000 words of jargon."
    )
    add_para(
        doc,
        "The product is NOT a spreadsheet dashboard, NOT a passive LMS, and NOT a generic quiz app. "
        "It is an exploration game where progress, mastery, conviction, and XP reward demonstrated understanding."
    )
    add_bullets(
        doc,
        [
            "Core metaphor: NVIDIA's 10-K as a world map with four research islands + a central Final Challenge hub.",
            "Primary demo company: NVIDIA (NVDA) — controlled MVP scope.",
            "Primary audience: Schools / classroom demos (Riverside Academy partner path in admin).",
            "Platform: Next.js 15 + React 19 web app; optional Capacitor wrappers for iOS/Android.",
            "Local dev: http://localhost:3003 (port is fixed — do not change without updating bookmarks and mobile sync)."
        ],
    )

    add_heading(doc, "2. Product Vision & Design North Star", 1)
    add_para(doc, "When a student opens Investor Quest, they should feel:", bold=True)
    add_bullets(
        doc,
        [
            "Curiosity — \"I want to explore this company.\"",
            "Progression — \"I'm leveling up as an investor.\"",
            "Achievement — \"I earned this badge / XP / island completion.\"",
            "Confidence — \"I understand how this business makes money and what risks matter.\""
        ],
    )
    add_para(doc, "Visual & interaction principles:", bold=True)
    add_bullets(
        doc,
        [
            "Premium fintech game aesthetic: dark sci-fi, glassmorphism, neon violet energy, cinematic glow.",
            "Important UI must stay clickable — decorative layers use pointer-events: none.",
            "Hydration-safe: no Math.random() during SSR; gate client-only logic in effects.",
            "Mobile-first responsiveness; full-desktop layouts for map previews and presenter mode.",
            "Frame actions as unlocking knowledge, building conviction, and leveling up — never \"reading reports.\""
        ],
    )

    add_heading(doc, "3. The Learning Journey (Schools Demo)", 1)
    add_para(doc, "Canonical presenter flow (scripted tour):", bold=True)
    add_numbered(
        doc,
        [
            "Logo Intro — brand entry, cinematic tone.",
            "Mission Brief Cards — welcome to NVIDIA's 10-K; explain the island quest and 10K Rookie Badge reward.",
            "Logo Reveal — product identity moment.",
            "Onboarding (Stocks Experience) — orientation.",
            "Map Brief — transition into the world map.",
            "Quest Map — choose an island (Business unlocked first).",
            "Business Island — first pillar hub and quests.",
            "Business Quest — read content cards + pass quizzes.",
            "Conviction — submit pillar conviction (unlocks progression signals).",
            "Profile — learner progress surface.",
            "XP Ladder — mastery / level visualization.",
            "Final Challenge — 10-K capstone quiz (12 questions, 3 per pillar).",
            "Map Return — completion loop back to world map."
        ],
    )
    add_para(doc, "Entry URL for live demo: http://localhost:3003/schools/demo")

    add_heading(doc, "4. World Map Structure", 1)
    add_para(doc, "All map variants share the same information architecture:", bold=True)
    add_bullets(
        doc,
        [
            "Center: 10K Hub — Final Challenge / spawn point.",
            "Top-left: Business — How do they make money?",
            "Top-right: Risks (Forces) — What are the threats?",
            "Bottom-left: Financial — Are they financially strong?",
            "Bottom-right: Management — Would you trust them?"
        ],
    )
    add_para(doc, "Map style previews (A–G) — same game state, different visual philosophy:", bold=True)
    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    headers = ["Label", "Route", "Style intent"]
    rows = [
        ("A — Cinematic", "/schools/map", "Production map — illustrated new-map.png, pillar cards"),
        ("B — Cartoon", "/schools/preview/map-cartoon", "Mario/Pokémon overworld, chunky regions"),
        ("C — Prodigy", "/schools/preview/map-prodigy", "Fantasy RPG educational realm"),
        ("D — DragonBox", "/schools/preview/map-dragonbox", "Minimal, Apple-like clarity"),
        ("E — Roblox", "/schools/preview/map-roblox", "Blocky adventure world, bright colors"),
        ("F — Khan Academy", "/schools/preview/map-khan", "LMS-style module cards, no world map"),
        ("G — Legends of Learning", "/schools/preview/map-legends", "Vibrant floating islands, edu-game art"),
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, (label, route, intent) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = label
        table.rows[r].cells[1].text = route
        table.rows[r].cells[2].text = intent
    doc.add_paragraph()

    add_heading(doc, "5. The Four Pillars (Research Islands)", 1)
    pillars = [
        ("Business", "What the company does, who it serves, and why it wins.", "Map the model: products, customers, revenue, advantage."),
        ("Forces (Risks)", "Competitive dynamics, industry structure, disruption.", "Decode rivals, suppliers, substitutes, new entrants."),
        ("Financials", "Cash flows, margins, capital allocation.", "Trace revenue drivers, balance sheet, capital redeployment."),
        ("Management", "Incentives, execution, decision quality.", "Judge track record, communication, stewardship."),
    ]
    for name, subtitle, desc in pillars:
        add_heading(doc, name, 2)
        add_para(doc, subtitle)
        add_para(doc, desc)

    add_heading(doc, "6. Quest & Content Model", 1)
    add_bullets(
        doc,
        [
            "Pillars are global; quest content is scoped per company (library in src/data/quests/).",
            "Each quest has content cards (reading) + quizzes (mastery proof).",
            "Quiz pass awards XP; \"Mark as Read\" tracks reading without XP.",
            "Island completion = all quests in pillar completed.",
            "Controlled demo (NVDA) exposes curated quest slugs per pillar — see controlledDemo.ts.",
            "Content pipeline supports AI generation with jargon rewrite loops (QUEST_FAST_DEV for local iteration).",
            "Quiz formats: multiple-choice, scenario, flashcard variants; copy standards in .cursor/rules/."
        ],
    )

    add_heading(doc, "7. Progression, XP & Badges", 1)
    add_para(doc, "Game state lives in a pure reducer engine (src/engine/). UI reads via selectors; never mutates state directly.", bold=True)
    add_bullets(
        doc,
        [
            "XP_SECTION_QUIZ: 50 · XP_QUIZ_PERFECT_BONUS: 25 · XP_ISLAND_COMPLETION: 200",
            "TEN_K_ROOKIE_CHALLENGE_XP: 1000 (final challenge)",
            "Levels derived from total XP (levelProgress selector).",
            "Badges: first-quest, first-pillar, all-pillars, quiz streaks, ten-k-rookie, etc.",
            "10K Rookie Badge: complete all four pillars + final challenge.",
            "Final Challenge unlock: all pillars complete + reading complete + conviction submitted per pillar.",
            "Persistence: localStorage key investor-quest::state (versioned migrations)."
        ],
    )

    add_heading(doc, "8. Key Features to Replicate in a Prototype", 1)
    add_heading(doc, "Must-have (MVP prototype)", 2)
    add_numbered(
        doc,
        [
            "World map with 4 islands + central 10K hub (pick one visual style or ship Cinematic A first).",
            "Locked / unlocked / in-progress / completed states per island.",
            "At least one fully playable pillar (Business) with 2–3 quest cards + quizzes.",
            "XP on quiz pass + island completion celebration.",
            "Mission brief onboarding (3–5 screens explaining the 10-K quest).",
            "Progress persistence (localStorage minimum).",
            "Mobile-friendly touch targets; desktop-first map layout."
        ],
    )
    add_heading(doc, "Should-have (parity with this repo)", 2)
    add_bullets(
        doc,
        [
            "Conviction submission flow per pillar.",
            "Final Challenge gate + 10K Rookie Badge.",
            "Profile + XP ladder screens.",
            "Presenter demo route (/schools/demo) with scripted step navigation.",
            "Admin quest generation + game health probes.",
            "Multiple map style experiments (preview routes B–G)."
        ],
    )
    add_heading(doc, "Nice-to-have", 2)
    add_bullets(
        doc,
        [
            "Capacitor mobile shells.",
            "Supabase analytics / remote progression sync.",
            "SEC filing ingestion pipeline.",
            "Multi-company support beyond NVDA."
        ],
    )

    add_heading(doc, "9. Technical Architecture", 1)
    add_bullets(
        doc,
        [
            "Framework: Next.js 15 App Router (src/app/).",
            "UI: React 19, Tailwind CSS 4, Framer Motion, custom CSS in globals.css.",
            "State: GameProvider + engine reducer (src/engine/progression/).",
            "Data: pillars, companies, quest catalogs (src/data/).",
            "Schools layer: src/lib/schools/, src/components/schools/.",
            "Demo mode: CONTROLLED_DEMO_MODE, NVIDIA-only funnel.",
            "Auth/DB: Supabase (optional for prototype).",
            "Admin: /admin/quests, /admin/game-health, quest generation studio."
        ],
    )

    add_heading(doc, "10. Repository Map (High Level)", 1)
    paths = [
        ("src/app/schools/", "Schools routes — map, pillars, demo, previews"),
        ("src/components/schools/", "Map screens, mission brief, logo flows"),
        ("src/components/QuestQuizPanel.tsx", "Core quiz UX"),
        ("src/engine/", "Progression engine public API"),
        ("src/data/quests/", "Quest definitions & templates"),
        ("src/lib/schools/", "Map configs, demo hrefs, mission copy"),
        ("src/app/globals.css", "Map preview styles (cartoon, prodigy, etc.)"),
        (".cursor/rules/", "Product & copy standards for AI-assisted dev"),
    ]
    for path, desc in paths:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{path} — ").bold = True
        p.add_run(desc)

    add_heading(doc, "11. Local Development Setup", 1)
    add_numbered(
        doc,
        [
            "Prerequisites: Node.js >= 20.9, npm.",
            "Clone repo, run npm install (requires approval in Cursor sandbox).",
            "Copy .env.local template if provided; set QUEST_FAST_DEV=true for fast quest iteration.",
            "Run: npm run dev — wait for ✓ Ready on port 3003.",
            "Open: http://localhost:3003/schools/demo",
            "Typecheck: npx tsc --noEmit · Lint: npm run lint",
            "Do NOT change dev port 3003 without updating docs, Capacitor sync, and bookmarks."
        ],
    )

    add_heading(doc, "12. Prototype Build Recommendations", 1)
    add_para(
        doc,
        "For a developer building a working prototype modeled after Investor Quest, recommended phasing:"
    )
    add_numbered(
        doc,
        [
            "Week 1 — Shell: Next.js app, GameProvider stub, one map screen (static), pillar routing.",
            "Week 2 — Business pillar: 3 quests with hardcoded content + quiz component + XP.",
            "Week 3 — Progression: unlock next island, map state visuals, mission brief.",
            "Week 4 — Polish: final challenge stub, profile, presenter demo path, one alternate map style.",
        ],
    )
    add_para(doc, "Success criteria for prototype demo:", bold=True)
    add_bullets(
        doc,
        [
            "A student can complete Business island quizzes and see XP + progress update.",
            "Map clearly shows which island is next; locked islands are obvious.",
            "Onboarding explains the 10-K quest in under 60 seconds.",
            "The experience feels like a game, not homework — within 3 seconds of seeing the map."
        ],
    )

    add_heading(doc, "13. Copy & Pedagogy Anchors", 1)
    add_bullets(
        doc,
        [
            "Business: How do they make money?",
            "Risks: What are the threats?",
            "Financial: Are they financially strong?",
            "Management: Would you trust them?",
            "Reward framing: 10K Rookie Badge — first step to becoming a smarter investor.",
            "Avoid jargon in student-facing copy; quiz copy standards enforced in repo rules."
        ],
    )

    add_heading(doc, "14. What Not to Build (Anti-patterns)", 1)
    add_bullets(
        doc,
        [
            "Flat spreadsheet-style company dashboards.",
            "White corporate SaaS cards on a gray admin layout.",
            "Unskippable walls of 10-K text without quests or checkpoints.",
            "Random XP for scrolling — XP rewards demonstrated mastery only.",
            "Overlapping floating labels on the map that don't belong to a region.",
            "Breaking hydration with random layout or browser-only measurements on first paint."
        ],
    )

    add_heading(doc, "15. Contacts & Next Steps", 1)
    add_para(doc, "For the developer starting the prototype:")
    add_numbered(
        doc,
        [
            "Run the Schools demo locally and complete Business island end-to-end.",
            "Review map style previews A–G in the Schools sidebar to understand visual range.",
            "Read src/engine/progression/reducer.ts and selectors.ts for state contracts.",
            "Pick one map style for v1 prototype (recommend A Cinematic or G Legends for game feel).",
            "Align with product owner on scope: NVDA-only vs multi-company, web-only vs mobile."
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("Generated from the Investor Quest codebase — investor-quest repository.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
